# SPDX-License-Identifier: PolyForm-Noncommercial-1.0.0
# Required Notice: Copyright © 2026 Jacob Iannotti. Commercial rights reserved. See LICENSE.md.
import os
import mlx_whisper
import torch
import subprocess
import soundfile as sf
import gc
import re
from pyannote.audio import Pipeline
from datetime import timedelta
from docx import Document

# CONFIGURATION
WATCH_FOLDER = "/Users/jacobiannotti/Desktop/DUMP/Transcription Folder- Walter Reed Project /transcribe_input"
OUTPUT_FOLDER = "/Users/jacobiannotti/Desktop/DUMP/Transcription Folder- Walter Reed Project /transcribe_output"
MODEL_REPO = "mlx-community/whisper-medium.en-mlx" 
EXTENSIONS = {".mp4", ".mkv", ".mov", ".webm", ".avi", ".m4a", ".mp3", ".wav"}

def format_timestamp(seconds):
    return str(timedelta(seconds=int(seconds)))

def get_clean_output_name(filename):
    match = re.search(r"(20\d{2})(\d{2})(\d{2})", filename)
    if match:
        clean_date = f"{match.group(1)}-{match.group(2)}-{match.group(3)}"
        clean_name = filename.replace(match.group(0), "")
        clean_name = re.sub(r"_\d{6}", "", clean_name)
        clean_name = clean_name.replace("_", " ").replace("-", " ").strip()
        clean_name = re.sub(r"\s+", " ", clean_name)
        return f"{clean_date} - {clean_name}"
    else:
        return os.path.splitext(filename)[0]

def extract_audio_wav(video_path):
    base_name = os.path.splitext(video_path)[0]
    wav_path = f"{base_name}_temp_audio.wav"
    command = ["ffmpeg", "-y", "-i", video_path, "-vn", "-acodec", "pcm_s16le", "-ar", "16000", "-ac", "1", wav_path]
    subprocess.run(command, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
    return wav_path if os.path.exists(wav_path) else None

def process_batch():
    hf_token = os.environ.get("HF_TOKEN")
    if not hf_token:
        print("❌ ERROR: HF_TOKEN not found.")
        return

    files = sorted([f for f in os.listdir(WATCH_FOLDER) if os.path.splitext(f)[1].lower() in EXTENSIONS and "_temp_audio" not in f])
    
    for filename in files:
        video_path = os.path.join(WATCH_FOLDER, filename)
        print(f"\n▶️  Processing: {filename}")
        
        try:
            temp_wav = extract_audio_wav(video_path)
            if not temp_wav: continue
            
            # --- STEP 1: TRANSCRIBE ---
            print("   Step 1: Transcribing (MLX)...")
            transcript_result = mlx_whisper.transcribe(temp_wav, path_or_hf_repo=MODEL_REPO, word_timestamps=True)
            
            # Clear Whisper from memory
            gc.collect()
            torch.mps.empty_cache()

            # --- STEP 2: DIARIZATION ---
            print("   Step 2: Identifying Speakers (MPS)...")
            pipeline = Pipeline.from_pretrained("pyannote/speaker-diarization-3.1", token=hf_token)
            pipeline.to(torch.device("mps"))
            
            waveform, sample_rate = sf.read(temp_wav)
            waveform_t = torch.from_numpy(waveform).float().unsqueeze(0).to(torch.device("mps"))
            
            raw_diarization = pipeline({"waveform": waveform_t, "sample_rate": sample_rate})
            
            # THE FIX: Safely unpack the diarization object
            diarization = getattr(raw_diarization, "speaker_diarization", raw_diarization)
            
            # Immediately free memory
            del pipeline
            del waveform_t
            gc.collect()
            torch.mps.empty_cache()

            # --- STEP 3: DOCUMENT ---
            print("   Step 3: Saving Document...")
            doc = Document()
            clean_title = get_clean_output_name(filename)
            doc.add_heading(clean_title, 0)
            
            speaker_turns = [{"start": t.start, "end": t.end, "speaker": s} for t, _, s in diarization.itertracks(yield_label=True)]
            current_speaker, buffer_text, current_start = None, [], 0
            
            for segment in transcript_result['segments']:
                s_start, s_end, text = segment['start'], segment['end'], segment['text'].strip()
                best_s = "Unknown"
                max_ov = 0
                for turn in speaker_turns:
                    ov = max(0, min(s_end, turn["end"]) - max(s_start, turn["start"]))
                    if ov > max_ov:
                        max_ov, best_s = ov, turn["speaker"]
                
                if best_s != current_speaker:
                    if current_speaker is not None:
                        p = doc.add_paragraph()
                        p.add_run(f"[{format_timestamp(current_start)}] {str(current_speaker).replace('SPEAKER_', 'Speaker ')}:").bold = True
                        p.add_run(f"\n{' '.join(buffer_text)}")
                    current_speaker, current_start, buffer_text = best_s, s_start, [text]
                else:
                    buffer_text.append(text)

            if buffer_text:
                p = doc.add_paragraph()
                p.add_run(f"[{format_timestamp(current_start)}] {str(current_speaker).replace('SPEAKER_', 'Speaker ')}:").bold = True
                p.add_run(f"\n{' '.join(buffer_text)}")

            doc.save(os.path.join(OUTPUT_FOLDER, f"{clean_title}.docx"))
            print(f"✅ Success.")
            
            os.remove(video_path)
            if os.path.exists(temp_wav):
                os.remove(temp_wav)

        except Exception as e:
            print(f"❌ Error: {e}")

if __name__ == "__main__":
    process_batch()
